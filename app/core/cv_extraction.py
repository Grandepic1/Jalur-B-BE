import re
import unicodedata
from io import BytesIO
from multiprocessing import get_context
from multiprocessing.connection import Connection
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

MAX_CV_BYTES = 5 * 1024 * 1024
MAX_CV_PAGES = 20
MAX_CV_TEXT_CHARS = 30_000
MIN_CV_TEXT_CHARS = 50
MAX_DOCX_FILES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
CV_EXTRACTION_TIMEOUT_SECONDS = 15
CV_EXTRACTION_MEMORY_BYTES = 256 * 1024 * 1024

PDF_CONTENT_TYPE = "application/pdf"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class CVExtractionError(ValueError):
    pass


def detect_cv_type(content: bytes, file_name: str) -> tuple[str, str] | None:
    extension = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""
    if extension == "pdf" and content.startswith(b"%PDF-"):
        return PDF_CONTENT_TYPE, ".pdf"
    if extension == "docx" and content.startswith(b"PK\x03\x04"):
        return DOCX_CONTENT_TYPE, ".docx"
    return None


def _normalize_text(value: str) -> str:
    value = unicodedata.normalize("NFKC", value).replace("\x00", "")
    lines = [re.sub(r"\s+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line)[:MAX_CV_TEXT_CHARS]


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content), strict=False)
        if reader.is_encrypted:
            raise CVExtractionError("Password-protected PDFs are not supported")
        if len(reader.pages) > MAX_CV_PAGES:
            raise CVExtractionError(f"PDFs may contain at most {MAX_CV_PAGES} pages")
        parts: list[str] = []
        extracted_chars = 0
        for page in reader.pages:
            page_text = page.extract_text() or ""
            remaining = MAX_CV_TEXT_CHARS - extracted_chars
            parts.append(page_text[:remaining])
            extracted_chars += min(len(page_text), remaining)
            if extracted_chars >= MAX_CV_TEXT_CHARS:
                break
        text = "\n".join(parts)
    except CVExtractionError:
        raise
    except Exception as exc:
        raise CVExtractionError("The PDF could not be read") from exc
    return _normalize_text(text)


def _extract_docx_text(content: bytes) -> str:
    try:
        with ZipFile(BytesIO(content)) as archive:
            entries = archive.infolist()
            if "word/document.xml" not in archive.namelist():
                raise CVExtractionError("The DOCX file is missing document content")
            if (
                len(entries) > MAX_DOCX_FILES
                or sum(entry.file_size for entry in entries)
                > MAX_DOCX_UNCOMPRESSED_BYTES
            ):
                raise CVExtractionError("The DOCX file expands beyond the safe limit")
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
    except CVExtractionError:
        raise
    except Exception as exc:
        raise CVExtractionError("The DOCX file could not be read") from exc
    return _normalize_text("\n".join(parts))


def extract_cv_text(content: bytes, content_type: str) -> str:
    if content_type == PDF_CONTENT_TYPE:
        text = _extract_pdf_text(content)
    elif content_type == DOCX_CONTENT_TYPE:
        text = _extract_docx_text(content)
    else:
        raise CVExtractionError("Unsupported CV file type")
    if len(text) < MIN_CV_TEXT_CHARS:
        raise CVExtractionError(
            "The CV contains too little readable text; scanned PDFs are not supported"
        )
    return text


def _extract_cv_text_worker(
    connection: Connection,
    content: bytes,
    content_type: str,
) -> None:
    try:
        try:
            import resource

            resource.setrlimit(
                resource.RLIMIT_AS,
                (CV_EXTRACTION_MEMORY_BYTES, CV_EXTRACTION_MEMORY_BYTES),
            )
            resource.setrlimit(
                resource.RLIMIT_CPU,
                (
                    CV_EXTRACTION_TIMEOUT_SECONDS,
                    CV_EXTRACTION_TIMEOUT_SECONDS,
                ),
            )
        except (ImportError, OSError, ValueError):
            pass
        connection.send((True, extract_cv_text(content, content_type)))
    except BaseException as exc:
        message = (
            str(exc)
            if isinstance(exc, CVExtractionError)
            else "The CV could not be read safely"
        )
        connection.send((False, message))
    finally:
        connection.close()


def extract_cv_text_isolated(content: bytes, content_type: str) -> str:
    context = get_context("spawn")
    receiving_connection, sending_connection = context.Pipe(duplex=False)
    process = context.Process(
        target=_extract_cv_text_worker,
        args=(sending_connection, content, content_type),
    )
    process.start()
    sending_connection.close()
    try:
        if not receiving_connection.poll(CV_EXTRACTION_TIMEOUT_SECONDS):
            raise CVExtractionError("CV text extraction timed out")
        try:
            succeeded, result = receiving_connection.recv()
        except EOFError as exc:
            raise CVExtractionError("The CV could not be read safely") from exc
    finally:
        receiving_connection.close()
        if process.is_alive():
            process.terminate()
        process.join()
    if not succeeded:
        raise CVExtractionError(result)
    return result
