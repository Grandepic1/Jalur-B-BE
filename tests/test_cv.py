from datetime import UTC, datetime, timedelta
from io import BytesIO
from types import SimpleNamespace
from unittest import IsolatedAsyncioTestCase, TestCase
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import UUID, uuid4

from docx import Document
from fastapi import HTTPException, UploadFile
from pydantic import ValidationError
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.cv import (
    _create_preview_token,
    _decode_preview_token,
    confirm_cv,
    preview_cv,
)
from app.core.config import settings
from app.core.cv_extraction import (
    DOCX_CONTENT_TYPE,
    MAX_CV_TEXT_CHARS,
    PDF_CONTENT_TYPE,
    CVExtractionError,
    detect_cv_type,
    extract_cv_text,
)
from app.models.cv import (
    CVConfirmRequest,
    CVConfirmResponse,
    CVExperience,
    CVExtractionAIResult,
    CVPreviewTokenData,
    UserCV,
)
from app.models.master import Skill
from app.models.profile import UserProfile


class CVExtractionTests(TestCase):
    def test_freelance_experience_can_omit_company_and_details(self) -> None:
        experience = CVExperience(
            role="Freelancer",
            company=None,
            start_date=None,
            end_date=None,
            description=None,
        )

        self.assertIsNone(experience.company)
        self.assertIsNone(experience.description)

    def test_profile_and_skills_are_normalized_for_preview(self) -> None:
        extraction = CVExtractionAIResult(
            profile={
                "full_name": "  Joan   Orlando ",
                "current_role_name": " Backend Engineer ",
                "industry_name": None,
                "work_duration_months": 24,
                "daily_activities": "  Membangun API dan menjaga layanan produksi. ",
            },
            skills=[" Python ", "python", " API Design "],
            experiences=[],
        )

        self.assertEqual(extraction.profile.full_name, "Joan Orlando")
        self.assertEqual(extraction.skills, ["Python", "API Design"])

    def test_preview_token_data_rejects_unknown_fields(self) -> None:
        with self.assertRaises(ValidationError):
            CVPreviewTokenData(
                preview_id=uuid4(),
                user_id=7,
                expires_at=datetime.now(UTC) + timedelta(hours=1),
                model="test-model",
                unknown="Untrusted override",
            )

    def test_confirm_request_accepts_edited_reviewed_values(self) -> None:
        request = CVConfirmRequest(
            preview_token="signed-token",
            profile={
                "full_name": "Joan Orlando",
                "current_role_name": "Backend Engineer",
                "industry_name": "Technology",
                "work_duration_months": 24,
                "daily_activities": "Membangun API dan menjaga layanan produksi.",
            },
            skills=[" Python ", "python", "API Design"],
            experiences=[
                {
                    "role": "Backend Engineer",
                    "company": "Contoh Digital",
                    "start_date": "2023",
                    "end_date": "Sekarang",
                    "description": "Membangun API dan menjaga layanan produksi.",
                }
            ],
        )

        self.assertEqual(request.profile.full_name, "Joan Orlando")
        self.assertEqual(request.skills, ["Python", "API Design"])
        self.assertEqual(request.experiences[0].company, "Contoh Digital")

    def test_confirm_request_rejects_unknown_fields(self) -> None:
        with self.assertRaises(ValidationError):
            CVConfirmRequest(
                preview_token="signed-token",
                profile={},
                skills=[],
                experiences=[],
                file_name="untrusted-override",
            )

    def test_confirm_request_enforces_skill_and_experience_caps(self) -> None:
        with self.assertRaises(ValidationError):
            CVConfirmRequest(
                preview_token="signed-token",
                profile={},
                skills=[f"Skill {index}" for index in range(21)],
                experiences=[],
            )
        with self.assertRaises(ValidationError):
            CVConfirmRequest(
                preview_token="signed-token",
                profile={},
                skills=[],
                experiences=[
                    {"role": f"Peran {index}", "description": "Bekerja rutin harian."}
                    for index in range(13)
                ],
            )

    def test_confirm_request_blank_profile_fields_stay_null(self) -> None:
        request = CVConfirmRequest(
            preview_token="signed-token",
            profile={"full_name": "   "},
            skills=[],
            experiences=[],
        )

        self.assertIsNone(request.profile.full_name)

    def test_experience_rejects_a_whitespace_only_role(self) -> None:
        with self.assertRaises(ValidationError):
            CVExperience(role="   ")

    def test_detects_supported_file_signatures(self) -> None:
        self.assertEqual(
            detect_cv_type(b"%PDF-1.7 content", "resume.PDF"),
            (PDF_CONTENT_TYPE, ".pdf"),
        )
        self.assertEqual(
            detect_cv_type(b"PK\x03\x04 content", "resume.docx"),
            (DOCX_CONTENT_TYPE, ".docx"),
        )
        self.assertIsNone(detect_cv_type(b"PK\x03\x04 content", "resume.doc"))
        self.assertIsNone(detect_cv_type(b"not a pdf", "resume.pdf"))

    def test_extracts_and_normalizes_docx_text(self) -> None:
        document = Document()
        document.add_heading("Riwayat Karier", level=1)
        document.add_paragraph(
            "Backend Engineer di Contoh Digital dari 2023 sampai sekarang."
        )
        document.add_paragraph(
            "Membangun API, menjaga layanan produksi, dan mempercepat deployment."
        )
        stream = BytesIO()
        document.save(stream)

        text = extract_cv_text(stream.getvalue(), DOCX_CONTENT_TYPE)

        self.assertIn("Backend Engineer di Contoh Digital", text)
        self.assertNotIn("  ", text)
        self.assertLessEqual(len(text), MAX_CV_TEXT_CHARS)

    def test_rejects_docx_without_meaningful_text(self) -> None:
        document = Document()
        document.add_paragraph("CV")
        stream = BytesIO()
        document.save(stream)

        with self.assertRaisesRegex(CVExtractionError, "too little readable text"):
            extract_cv_text(stream.getvalue(), DOCX_CONTENT_TYPE)

    def test_rejects_docx_that_expands_beyond_limit(self) -> None:
        document = Document()
        document.add_paragraph(
            "Backend Engineer dengan pengalaman membangun dan menjaga layanan produksi."
        )
        stream = BytesIO()
        document.save(stream)

        with (
            patch("app.core.cv_extraction.MAX_DOCX_UNCOMPRESSED_BYTES", 1),
            self.assertRaisesRegex(CVExtractionError, "safe limit"),
        ):
            extract_cv_text(stream.getvalue(), DOCX_CONTENT_TYPE)


class CVConfirmationTests(IsolatedAsyncioTestCase):
    def setUp(self) -> None:
        super().setUp()
        secret_patch = patch.object(settings, "jwt_secret_key", "x" * 64)
        secret_patch.start()
        self.addCleanup(secret_patch.stop)

    def _signed_token(self, preview_id: UUID | None = None) -> str:
        preview = CVPreviewTokenData(
            preview_id=preview_id or uuid4(),
            user_id=7,
            expires_at=datetime.now(UTC) + timedelta(hours=1),
            model="test-model",
        )
        return _create_preview_token(preview)

    async def _confirm(
        self, request: CVConfirmRequest, db: MagicMock
    ) -> CVConfirmResponse:
        return await confirm_cv(request, SimpleNamespace(id=7), db)

    def _profile(self, now: datetime) -> UserProfile:
        return UserProfile(
            id=11,
            user_id=7,
            full_name="Old Name",
            current_role_name="Old Role",
            industry_name="Old Industry",
            work_duration_months=1,
            is_first_job=False,
            daily_activities="Old activities",
            career_goal=None,
            target_role_name=None,
            target_industry_name=None,
            onboarding_completed_at=now,
            created_at=now,
            updated_at=now,
        )

    def _db(self) -> MagicMock:
        db = MagicMock(spec=AsyncSession)
        db.execute = AsyncMock()
        db.delete = AsyncMock()
        db.commit = AsyncMock()
        db.rollback = AsyncMock()
        db.refresh = AsyncMock()
        db.flush = AsyncMock()
        return db

    async def test_preview_returns_signed_result_without_persisting(self) -> None:
        document = Document()
        document.add_paragraph(
            "Joan Orlando, Backend Engineer dengan pengalaman membangun API produksi."
        )
        stream = BytesIO()
        document.save(stream)
        upload = UploadFile(filename="resume.docx", file=BytesIO(stream.getvalue()))
        extraction = CVExtractionAIResult(
            profile={
                "full_name": "Joan Orlando",
                "current_role_name": "Backend Engineer",
                "industry_name": None,
                "work_duration_months": None,
                "daily_activities": "Membangun dan menjaga API produksi.",
            },
            skills=["API Design"],
            experiences=[],
        )
        provider = MagicMock(model="test-model")
        provider.generate_structured = AsyncMock(return_value=extraction)
        db = MagicMock(spec=AsyncSession)
        db.scalar = AsyncMock(side_effect=[11])
        db.rollback = AsyncMock()

        with patch.object(settings, "jwt_secret_key", "x" * 64):
            response = await preview_cv(
                upload,
                SimpleNamespace(id=7),
                db,
                provider,
            )

        with patch.object(settings, "jwt_secret_key", "x" * 64):
            token_data = _decode_preview_token(response.preview_token)
        self.assertEqual(response.profile.full_name, "Joan Orlando")
        self.assertEqual(response.skills, ["API Design"])
        self.assertEqual(token_data.preview_id, response.preview_id)
        self.assertEqual(token_data.user_id, 7)
        self.assertEqual(token_data.model, "test-model")
        db.add.assert_not_called()

    async def test_confirmation_applies_edited_payload_atomically(self) -> None:
        now = datetime.now(UTC)
        token = self._signed_token()
        profile = self._profile(now)
        skill = Skill(
            id=5,
            name="Python",
            category="technical",
            market_trend="stable",
        )
        db = self._db()
        db.scalar = AsyncMock(
            side_effect=[None, 7, None, profile, None]
        )
        db.scalars = AsyncMock(
            side_effect=[
                SimpleNamespace(all=lambda: [skill]),
                SimpleNamespace(all=lambda: [skill]),
            ]
        )
        preview_id = _decode_preview_token(token).preview_id

        response = await self._confirm(
            CVConfirmRequest(
                preview_token=token,
                profile={
                    "full_name": "Joan Orlando",
                    "current_role_name": "Backend Engineer",
                    "industry_name": "Technology",
                    "work_duration_months": 24,
                    "daily_activities": "Membangun API dan menjaga layanan produksi.",
                },
                skills=["Python"],
                experiences=[
                    {
                        "role": "Backend Engineer",
                        "company": "Contoh Digital",
                        "start_date": "2023",
                        "end_date": "Sekarang",
                        "description": "Membangun API dan menjaga layanan produksi.",
                    }
                ],
            ),
            db,
        )

        self.assertEqual(profile.full_name, "Joan Orlando")
        self.assertEqual(profile.current_role_name, "Backend Engineer")
        self.assertEqual(profile.industry_name, "Technology")
        self.assertEqual(response.cv.experiences[0].role, "Backend Engineer")
        self.assertEqual(response.cv.model, "test-model")
        self.assertEqual([item.name for item in response.skills], ["Python"])
        saved_cv = next(
            call.args[0]
            for call in db.add.call_args_list
            if isinstance(call.args[0], UserCV)
        )
        self.assertIsInstance(saved_cv, UserCV)
        self.assertEqual(saved_cv.source_preview_id, preview_id)
        self.assertIsNone(saved_cv.file_name)
        self.assertIsNone(saved_cv.storage_object_path)
        self.assertEqual(
            saved_cv.experiences[0]["role"], "Backend Engineer"
        )
        self.assertEqual(db.execute.await_count, 2)
        self.assertEqual(db.commit.await_count, 1)

    async def test_confirmation_null_fields_keep_existing_values(self) -> None:
        now = datetime.now(UTC)
        profile = self._profile(now)
        db = self._db()
        db.scalar = AsyncMock(
            side_effect=[None, 7, None, profile, None]
        )
        db.scalars = AsyncMock(side_effect=[SimpleNamespace(all=lambda: [])])

        await self._confirm(
            CVConfirmRequest(
                preview_token=self._signed_token(),
                profile={"full_name": None, "industry_name": None},
                skills=[],
                experiences=[],
            ),
            db,
        )

        self.assertEqual(profile.full_name, "Old Name")
        self.assertEqual(profile.industry_name, "Old Industry")
        self.assertEqual(db.commit.await_count, 1)

    async def test_confirmation_merges_skills_without_deleting_existing(self) -> None:
        now = datetime.now(UTC)
        profile = self._profile(now)
        python = Skill(
            id=5,
            name="Python",
            category="technical",
            market_trend="stable",
        )
        db = self._db()
        db.scalar = AsyncMock(side_effect=[None, 7, None, profile, None])
        db.scalars = AsyncMock(
            side_effect=[
                SimpleNamespace(all=lambda: [python]),
                SimpleNamespace(all=lambda: [python]),
            ]
        )

        await self._confirm(
            CVConfirmRequest(
                preview_token=self._signed_token(),
                profile={},
                skills=["Python", "API Design"],
                experiences=[],
            ),
            db,
        )

        self.assertEqual(db.execute.await_count, 2)
        self.assertEqual(db.commit.await_count, 1)

    async def test_confirmation_retry_returns_the_already_applied_result(self) -> None:
        now = datetime.now(UTC)
        preview_id = uuid4()
        token = self._signed_token(preview_id)
        cv = UserCV(
            id=3,
            user_id=7,
            file_name=None,
            file_size=None,
            content_type=None,
            storage_object_path=None,
            source_preview_id=preview_id,
            experiences=[],
            provider_model="test-model",
            uploaded_at=now,
        )
        profile = self._profile(now)
        db = self._db()
        db.scalar = AsyncMock(side_effect=[SimpleNamespace(), cv, profile])
        db.scalars = AsyncMock(side_effect=[SimpleNamespace(all=lambda: [])])

        response = await self._confirm(
            CVConfirmRequest(
                preview_token=token,
                profile={"full_name": "Edited Later"},
                skills=[],
                experiences=[],
            ),
            db,
        )

        self.assertEqual(response.cv.model, "test-model")
        self.assertEqual(response.profile.full_name, "Old Name")
        db.commit.assert_not_called()

    async def test_consumed_superseded_preview_cannot_be_replayed(self) -> None:
        now = datetime.now(UTC)
        preview_id = uuid4()
        token = self._signed_token(preview_id)
        current_cv = UserCV(
            id=3,
            user_id=7,
            file_name=None,
            file_size=None,
            content_type=None,
            storage_object_path=None,
            source_preview_id=uuid4(),
            experiences=[],
            provider_model="test-model",
            uploaded_at=now,
        )
        db = self._db()
        db.scalar = AsyncMock(side_effect=[SimpleNamespace(), current_cv])

        with self.assertRaises(HTTPException) as raised:
            await confirm_cv(
                CVConfirmRequest(
                    preview_token=token,
                    profile={},
                    skills=[],
                    experiences=[],
                ),
                SimpleNamespace(id=7),
                db,
            )

        self.assertEqual(raised.exception.status_code, 409)
        db.commit.assert_not_called()

    async def test_replacing_legacy_stored_cv_enqueues_old_file_cleanup(self) -> None:
        now = datetime.now(UTC)
        preview_id = uuid4()
        token = self._signed_token(preview_id)
        profile = self._profile(now)
        legacy_cv = UserCV(
            id=3,
            user_id=7,
            file_name="old-resume.pdf",
            file_size=100,
            content_type=PDF_CONTENT_TYPE,
            storage_object_path="users/7/cv/legacy.pdf",
            source_preview_id=uuid4(),
            experiences=[],
            provider_model="test-model",
            uploaded_at=now,
        )
        db = self._db()
        db.scalar = AsyncMock(side_effect=[None, 7, None, profile, legacy_cv])
        db.scalars = AsyncMock(side_effect=[SimpleNamespace(all=lambda: [])])
        immediate_cleanup = AsyncMock()

        with patch(
            "app.api.cv.process_storage_deletion_path",
            new=immediate_cleanup,
        ):
            await confirm_cv(
                CVConfirmRequest(
                    preview_token=token,
                    profile={},
                    skills=[],
                    experiences=[],
                ),
                SimpleNamespace(id=7),
                db,
            )

        self.assertEqual(legacy_cv.source_preview_id, preview_id)
        self.assertIsNone(legacy_cv.storage_object_path)
        self.assertIsNone(legacy_cv.file_name)
        enqueue_calls = [
            call.args[0]
            for call in db.execute.call_args_list
            if call.args
            and getattr(call.args[0], "table", None) is not None
            and call.args[0].table.name == "storage_deletion_jobs"
        ]
        self.assertEqual(len(enqueue_calls), 1)
        self.assertEqual(db.commit.await_count, 1)
        immediate_cleanup.assert_awaited_once_with(db, "users/7/cv/legacy.pdf")
